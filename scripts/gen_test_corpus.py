"""Generate a realistic synthetic subrogation test corpus under data/sample_uploads/.

Fictional names/values only (no real PII). Document formats mirror real artifacts
(CHP 555 collision report, CCC ONE repair estimate, Bosch CDR readout, ACORD 2
FNOL, insurer subrogation demand letter, recorded statements). Output spans all
four ingestion paths — PDF, DOCX, HTML, plain text/markdown — across three cases
designed to drive the three pipeline outcomes (pursue / decline / escalate).

    python scripts/gen_test_corpus.py

Re-runnable: overwrites the corpus each time.
"""
from __future__ import annotations

from pathlib import Path

from fpdf import FPDF
from docx import Document

BASE = Path(__file__).resolve().parent.parent / "data" / "sample_uploads"

_REPL = {
    "–": "-", "—": "-", "‘": "'", "’": "'",
    "“": '"', "”": '"', "§": "Sec.", "•": "*", "…": "...",
}


def _ascii(s: str) -> str:
    for k, v in _REPL.items():
        s = s.replace(k, v)
    return s.encode("latin-1", "replace").decode("latin-1")


def write_txt(path: Path, text: str) -> None:
    path.write_text(text.strip() + "\n", encoding="utf-8")


def write_html(path: Path, title: str, body_html: str) -> None:
    path.write_text(
        f"<!doctype html><html><head><meta charset='utf-8'><title>{title}</title></head>"
        f"<body>\n{body_html}\n</body></html>\n",
        encoding="utf-8",
    )


def write_docx(path: Path, blocks: list[tuple[str, str]]) -> None:
    doc = Document()
    for kind, text in blocks:
        if kind == "h1":
            doc.add_heading(text, level=1)
        elif kind == "h2":
            doc.add_heading(text, level=2)
        else:
            doc.add_paragraph(text)
    doc.save(str(path))


def write_pdf(path: Path, blocks: list[tuple[str, str]]) -> None:
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    for kind, text in blocks:
        if kind == "h1":
            pdf.set_font("Helvetica", "B", 14)
            pdf.multi_cell(0, 8, _ascii(text)); pdf.ln(2)
        elif kind == "h2":
            pdf.set_font("Helvetica", "B", 11)
            pdf.multi_cell(0, 7, _ascii(text)); pdf.ln(1)
        else:
            pdf.set_font("Helvetica", size=10)
            pdf.multi_cell(0, 5, _ascii(text)); pdf.ln(1)
    pdf.output(str(path))


# =============================================================================
# CASE 1 — PURSUE: other driver ran a red light; our insured had the green.
#   Insured: Daniel Cho (2023 Subaru Outback) | Other: Robert Hale (2019 Ram 1500)
#   Loss: Wilshire Blvd & Vermont Ave, Los Angeles, CA | Damages ~ $18,470 | CVC 21453
# =============================================================================
def case_pursue(d: Path) -> None:
    d.mkdir(parents=True, exist_ok=True)

    write_pdf(d / "collision_report_LA-2026-31882.pdf", [
        ("h1", "STATE OF CALIFORNIA - TRAFFIC COLLISION REPORT (CHP 555)"),
        ("p", "Report Number: LA-2026-31882    Investigating Agency: Los Angeles Police Department, Central Traffic Division"),
        ("p", "Date of Collision: 02/11/2026   Time: 0814 hours   NCIC: 1942"),
        ("p", "Location: Wilshire Blvd at Vermont Ave, Los Angeles, CA 90010"),
        ("h2", "PARTIES"),
        ("p", "Party 1 (P-1) DRIVER: Daniel Cho. Vehicle V-1: 2023 Subaru Outback, plate 8WJK221. Direction of travel: Northbound on Vermont Ave. Owner: Daniel Cho."),
        ("p", "Party 2 (P-2) DRIVER: Robert Hale. Vehicle V-2: 2019 RAM 1500, plate 5TRN908. Direction of travel: Eastbound on Wilshire Blvd. Owner: Robert Hale."),
        ("h2", "PRIMARY COLLISION FACTOR"),
        ("p", "VC Section Violated: 21453(a) - Party 2. Type of Collision: Broadside. Movement Preceding Collision (P-2): Proceeding Straight. Movement Preceding Collision (P-1): Proceeding Straight."),
        ("h2", "STATEMENTS"),
        ("p", "P-1 (Cho) stated: 'I had a green light going north on Vermont. As I entered the intersection the pickup came through from my left and hit my driver door. I never had time to brake.'"),
        ("p", "P-2 (Hale) stated: 'I thought the light was yellow. I might have entered late.'"),
        ("p", "Independent witness W-1 (Sandra Pierce), pedestrian at the northwest corner, stated: 'The Subaru clearly had the green light. The truck ran the red - it didn't even slow down.'"),
        ("h2", "AREA OF IMPACT"),
        ("p", "AOI was determined by physical evidence (gouge marks and debris field) to be within the number 2 northbound lane, approximately 18 feet north of the south crosswalk line."),
        ("h2", "OPINIONS AND CONCLUSIONS"),
        ("p", "Based on the physical evidence, the independent witness statement, and the signal timing data obtained from LADOT, it is my opinion that Party 2 (Hale) entered the intersection against a steady red signal in violation of CVC 21453(a) and was the sole proximate cause of this collision. Party 1 was traveling within the posted 35 mph limit and had the right of way. Reporting Officer: M. Delgado, Serial 30418."),
    ])

    write_txt(d / "repair_estimate_CCC.txt", """
CCC ONE ESTIMATE - Preliminary
Shop: Westgate Collision Center, 1422 S Vermont Ave, Los Angeles, CA
Claimant Vehicle: 2023 Subaru Outback Limited   VIN: 4S4BTGLD2P3xxxxxx   Mileage: 21,403
Insured: Daniel Cho     Claim: PA-2026-77310     Estimator: J. Whitfield

Line  Oper  Description                          Part Number     Qty   Price$    Labor
1     Repl  Door shell, front left               61021AL00A      1     842.00    3.5
2     R&I   Door trim panel, front left                          1               1.2
3     Repl  Outer mirror, left (power, heated)   91036AL21A      1     388.00    0.6
4     Repl  Front door glass, left               61011AL01A      1     214.00    1.0
5     Rpr   Quarter panel, left                                  1               4.8
6     Repl  Alloy wheel, front left              28111AL02A      1     596.00    0.5
7     R&R   Front suspension lower arm, left     20202AL09A      1     472.00    2.6
8     Refn  Door shell, front left (clearcoat)                   1               2.8
9     Blnd  Quarter panel, left (blend)                          1               1.5
10    Subl  Wheel alignment (4-wheel)                            1     149.00    T

TOTALS
Parts ................................. 3,253.00
Body Labor   22.5 hrs @ $62.00/hr ..... 1,395.00
Paint Labor   8.4 hrs @ $62.00/hr .....   520.80
Mechanical    5.2 hrs @ $98.00/hr .....   509.60
Paint Supplies ........................   294.00
Sublet ................................   149.00
Subtotal .............................. 6,121.40
Sales Tax @ 9.5% ......................   581.53
TOTAL COST OF REPAIRS ................. 6,702.93

Note: vehicle also declared a partial total on prior estimate; supplement pending.
Total documented damages for this claim, including diminished value and rental: $18,470.00
""")

    write_txt(d / "cdr_readout_V2_ram.txt", """
BOSCH CRASH DATA RETRIEVAL (CDR) REPORT
Vehicle: 2019 RAM 1500  (Party 2 - Robert Hale)
Imaged by: ACTAR #2841   File complete: Yes   Multi-event: 1 of 1

PRE-CRASH DATA (5 seconds, sampled at 1.0 s)
Time(s)   Vehicle Speed(mph)   Engine RPM   Throttle(% full)   Service Brake
-5.0      41                   1640         18.0               Off
-4.0      42                   1690         19.2               Off
-3.0      43                   1710         21.0               Off
-2.0      43                   1705         20.1               Off
-1.0      42                   1660         15.0               Off
 0.0      40                   1520          0.0               On

EVENT DATA
Maximum Recorded Velocity Change (Delta-V), Longitudinal: -16.4 MPH
Time of Maximum Delta-V: 86 ms
Driver Belt Switch Circuit: BUCKLED
Frontal Air Bag Deployment, Time to Deploy: 41 ms
Ignition Cycle, Crash: 9,517

Interpretation: Party 2's vehicle was traveling 40-43 mph in the 5 seconds before impact
with no brake application until 0.0 s, indicating the driver did not slow for the signal.
""")

    write_docx(d / "subrogation_demand_letter.docx", [
        ("p", "PACIFIC GUARDIAN INSURANCE - Subrogation Department"),
        ("p", "700 S Flower St, Suite 1200, Los Angeles, CA 90017"),
        ("p", "Date: March 3, 2026"),
        ("p", "TO: Claims Department, Frontier Mutual Casualty Co."),
        ("p", "RE: Our Insured (Daniel Cho) v. Your Insured (Robert Hale)"),
        ("p", "Our Claim No.: PA-2026-77310    Your Claim No.: FM-2026-44219    Policy No.: PG-AUT-5582013"),
        ("p", "Date of Loss: February 11, 2026    Loss Location: Wilshire Blvd at Vermont Ave, Los Angeles, CA"),
        ("h2", "Statement of Facts"),
        ("p", "On February 11, 2026, our insured Daniel Cho was traveling northbound on Vermont Avenue with a steady green signal when your insured, Robert Hale, entered the intersection eastbound against a steady red signal and struck the driver side of our insured's vehicle. The Los Angeles Police Department report (LA-2026-31882) cites your insured under CVC 21453(a) as the primary collision factor."),
        ("h2", "Liability"),
        ("p", "Your insured was negligent in the operation of his vehicle and is legally liable for the resulting damages. The independent witness statement, the area-of-impact evidence, and your insured's own event data recorder (showing 40-43 mph with no braking before impact) establish that your insured failed to stop for a steady red light. Liability rests entirely with your insured."),
        ("h2", "Damages"),
        ("p", "Vehicle repairs: $6,702.93. Diminished value: $4,100.00. Rental / loss of use (18 days): $1,260.00. Towing and storage: $407.00. Medical (insured cervical strain, see attached): $5,800.00. Deductible: $200.00. Total documented damages: $18,470.00."),
        ("h2", "Demand"),
        ("p", "We hereby demand reimbursement in the amount of $18,470.00, including our insured's deductible. Please remit payment within 30 days of the date of this letter. If we do not receive your response, this matter will be submitted to Arbitration Forums, Inc."),
    ])


# =============================================================================
# CASE 2 — DECLINE: our insured rear-ended a stopped vehicle (our fault).
#   Insured: Megan Ross (2022 Toyota RAV4) | Other: Luis Ortega (2020 Honda Civic)
#   Loss: I-210 W at Lake Ave offramp, Pasadena, CA | Damages ~ $3,150 | CVC 21703
# =============================================================================
def case_decline(d: Path) -> None:
    d.mkdir(parents=True, exist_ok=True)

    write_html(d / "police_report.html", "Traffic Collision Report PAS-2026-10744", """
<h1>Pasadena Police Department - Traffic Collision Report</h1>
<p><b>Report Number:</b> PAS-2026-10744 &nbsp; <b>Date:</b> 01/29/2026 &nbsp; <b>Time:</b> 1737 hours</p>
<p><b>Location:</b> Interstate 210 West at the Lake Avenue off-ramp, Pasadena, CA</p>
<h2>Parties</h2>
<p><b>Party 1 (Driver):</b> Megan Ross, 2022 Toyota RAV4, plate 9KPD774, traveling west on the I-210 off-ramp. <b>Owner:</b> Megan Ross.</p>
<p><b>Party 2 (Driver):</b> Luis Ortega, 2020 Honda Civic, plate 7BNM330, stopped at the bottom of the off-ramp for a red signal. <b>Owner:</b> Luis Ortega.</p>
<h2>Primary Collision Factor</h2>
<p>VC Section Violated: 21703 - Party 1 (following too closely). Type of Collision: Rear End.</p>
<h2>Statements</h2>
<p>Party 1 (Ross) stated: "Traffic stopped suddenly at the bottom of the ramp and I couldn't stop in time. I rear-ended the car in front of me. It was my fault."</p>
<p>Party 2 (Ortega) stated: "I was completely stopped at the red light when I was hit from behind."</p>
<h2>Opinions and Conclusions</h2>
<p>Physical evidence is consistent with the statements. Party 1 was following too closely and failed to stop for stopped traffic ahead, in violation of CVC 21703, and is the primary cause of this collision. Party 2 was lawfully stopped and bears no fault. Reporting Officer: K. Underwood, Serial 8842.</p>
""")

    write_html(d / "fnol_acord2.html", "ACORD 2 Automobile Loss Notice", """
<h1>AUTOMOBILE LOSS NOTICE (ACORD 2)</h1>
<p><b>Date of Loss:</b> 01/29/2026 1737 &nbsp; <b>Policy Number:</b> SS-AUT-9920431</p>
<p><b>Insured:</b> Megan Ross &nbsp; <b>Carrier:</b> Summit Shield Insurance</p>
<p><b>Description of Location:</b> I-210 West at Lake Ave off-ramp, Pasadena, CA</p>
<p><b>Description of Accident:</b> Insured was descending the off-ramp when traffic stopped for a red signal. Insured was unable to stop in time and struck the rear of the vehicle ahead (a 2020 Honda Civic that was fully stopped).</p>
<p><b>Police Department:</b> Pasadena PD &nbsp; <b>Report Number:</b> PAS-2026-10744</p>
<p><b>Insured Vehicle:</b> 2022 Toyota RAV4, VIN 2T3xxxxxxxNW00000, Plate 9KPD774, Owner Megan Ross. Driver License: CA D1184220.</p>
<p><b>Other Vehicle:</b> 2020 Honda Civic, Plate 7BNM330, Driver/Owner Luis Ortega.</p>
<p><b>Injured:</b> None reported. <b>Witness:</b> None.</p>
""")

    write_txt(d / "repair_estimate_insured_vehicle.txt", """
REPAIR ESTIMATE (Mitchell)
Shop: Crown City Auto Body, Pasadena, CA
Vehicle: 2022 Toyota RAV4 XLE (Insured - Megan Ross)   Claim: SS-2026-30021

Line  Oper  Description                       Qty   Price$   Labor
1     Repl  Front bumper cover                1     412.00   2.0
2     Repl  Bumper absorber                   1      96.00   0.4
3     Rpr   Hood                              1               2.5
4     Repl  Grille assembly                   1     268.00   0.7
5     Refn  Front bumper cover                1               2.2

TOTALS
Parts ............................. 776.00
Body Labor 5.6 hrs @ $60/hr ....... 336.00
Paint Labor 2.2 hrs @ $60/hr ...... 132.00
Paint Supplies .................... 110.00
Subtotal ......................... 1,354.00
Sales Tax @ 9.5% ..................  128.63
TOTAL ............................ 1,482.63

NOTE: This estimate is for OUR INSURED's own vehicle (the at-fault, following vehicle).
Other party (Ortega) rear damage estimate, provided separately: $3,150.00.
""")


# =============================================================================
# CASE 3 — ESCALATE: disputed left-turn vs. alleged speeding (~50/50, conflict).
#   Insured: Aisha Khan (2021 Mazda CX-5) | Other: Tyler Brooks (2018 Ford Mustang)
#   Loss: 4th St & Alameda Ave, Burbank, CA | Damages ~ $12,900 | CVC 21801 vs 22350
# =============================================================================
def case_escalate(d: Path) -> None:
    d.mkdir(parents=True, exist_ok=True)

    write_pdf(d / "collision_report_BUR-2026-5521.pdf", [
        ("h1", "BURBANK POLICE DEPARTMENT - TRAFFIC COLLISION REPORT"),
        ("p", "Report Number: BUR-2026-5521   Date: 03/19/2026   Time: 2106 hours"),
        ("p", "Location: 4th Street at Alameda Avenue, Burbank, CA"),
        ("h2", "PARTIES"),
        ("p", "Party 1 (Driver): Aisha Khan, 2021 Mazda CX-5, plate 6HTL451, attempting a left turn from eastbound 4th St onto northbound Alameda Ave."),
        ("p", "Party 2 (Driver): Tyler Brooks, 2018 Ford Mustang GT, plate 4RPW889, traveling westbound on 4th St (oncoming to P-1)."),
        ("h2", "STATEMENTS"),
        ("p", "P-1 (Khan) stated: 'I had a green light and waited to turn left. The oncoming lane looked clear, then the Mustang appeared very fast out of nowhere. He must have been speeding - I never would have turned otherwise.'"),
        ("p", "P-2 (Brooks) stated: 'I had the green and the right of way going straight. She turned left right in front of me. I was doing the limit, maybe a little over.'"),
        ("p", "Witness W-1 (no independent witness located). The intersection has no traffic camera. Signal timing confirms both directions of 4th St had a circular green; there is no protected left-turn arrow."),
        ("h2", "OTHER ASSOCIATED FACTORS"),
        ("p", "Roadway was dry; lighting was dark with street lights. The posted speed limit on 4th Street is 35 mph. Skid analysis was inconclusive due to ABS. Estimated speed of P-2 from crush and roadway evidence ranged 38-52 mph - the range is too wide to establish a definitive speed."),
        ("h2", "OPINIONS AND CONCLUSIONS"),
        ("p", "This collision presents conflicting accounts. Party 1 may have failed to yield to oncoming traffic when making a left turn (CVC 21801). Party 2 may have been exceeding the basic speed law (CVC 22350). The available physical evidence does not conclusively establish which factor was primary; fault is contested. Reporting Officer: R. Avila, Serial 6620."),
    ])

    write_docx(d / "recorded_statement_insured_khan.docx", [
        ("p", "RECORDED STATEMENT - SUMMARY"),
        ("p", "Adjuster: This is Brian Nolan with Meridian Auto Insurance taking a recorded statement of Aisha Khan, claim number ME-2026-66120, on March 22, 2026 at 10:15 a.m. Aisha, do I have your permission to record this conversation?"),
        ("p", "Khan: Yes."),
        ("p", "Adjuster: Tell me what happened."),
        ("p", "Khan: I was eastbound on 4th and I had a green light to turn left onto Alameda. I edged into the intersection and waited. The oncoming lane was clear when I looked. I started my turn and the Mustang came at me extremely fast - he had to be well over the limit. He hit my passenger side."),
        ("p", "Adjuster: How fast do you think he was going?"),
        ("p", "Khan: At least 50. It happened in a second. A normal-speed car I would have easily cleared."),
        ("p", "Adjuster: Is everything you've told me true and correct to the best of your knowledge?"),
        ("p", "Khan: Yes, it is."),
    ])

    write_txt(d / "recorded_statement_other_brooks.txt", """
RECORDED STATEMENT - SUMMARY
Adjuster: This is Dana Pruitt with Frontier Mutual Casualty taking a recorded statement of
Tyler Brooks, claim number FM-2026-50887, on March 23, 2026. Tyler, do I have your
permission to record?
Brooks: Yeah, that's fine.
Adjuster: Describe the accident.
Brooks: I was westbound on 4th with a green light, going straight. I had the right of way.
This SUV turned left directly across my path. I couldn't avoid it.
Adjuster: How fast were you going?
Brooks: The speed limit is 35. I was right around there, maybe 38, 40 tops. I was not speeding
the way she's claiming.
Adjuster: Did you brake?
Brooks: I slammed the brakes but there was no time. She turned right in front of me.
Adjuster: Is everything true and correct to the best of your knowledge?
Brooks: Yes.

NOTE: Damages to insured vehicle (Mazda CX-5), per estimate, total $12,900.00.
The two accounts conflict on speed and right-of-way; no independent witness or camera exists.
""")


def write_readme(base: Path) -> None:
    (base / "README.md").write_text("""# Sample Upload Corpus (synthetic, for stress testing)

Fictional documents (no real PII) modeled on real subrogation artifacts: CHP 555
collision reports, CCC ONE / Mitchell repair estimates, Bosch CDR readouts,
ACORD 2 FNOL, insurer subrogation demand letters, and recorded statements.
Regenerate with `python scripts/gen_test_corpus.py`.

Each folder is one case. Create a case in the UI, fill the suggested fields, and
upload **all** files in the folder. Formats span every ingestion path: PDF, DOCX,
HTML, plain text.

| Folder | Scenario | Suggested case fields (insured / other / jurisdiction / damages) | Expected outcome |
|---|---|---|---|
| `case_pursue_redlight/` | Other driver ran a red light; our insured had green. Clear other-party fault, high value. | Daniel Cho / Robert Hale / CA / **18470** | **pursue or escalate** (recovery >= $25k threshold escalates) |
| `case_decline_rearend/` | OUR insured rear-ended a stopped car (following too closely). We are at fault. | Megan Ross / Luis Ortega / CA / **3150** | **decline** (low recovery, our fault) |
| `case_escalate_leftturn/` | Disputed left-turn vs. alleged speeding. Conflicting accounts, ~50/50, no witness/camera. | Aisha Khan / Tyler Brooks / CA / **12900** | **escalate** (near 50/50, adjudicator disagreement) |

Files per case:

- **case_pursue_redlight/**: `collision_report_LA-2026-31882.pdf` (multi-page CHP 555),
  `repair_estimate_CCC.txt`, `cdr_readout_V2_ram.txt`, `subrogation_demand_letter.docx`
- **case_decline_rearend/**: `police_report.html`, `fnol_acord2.html`, `repair_estimate_insured_vehicle.txt`
- **case_escalate_leftturn/**: `collision_report_BUR-2026-5521.pdf`, `recorded_statement_insured_khan.docx`, `recorded_statement_other_brooks.txt`

Single-file quick test: `../sample_uploads/police_report_demo.txt` (one clean case).

All scenarios reference real California Vehicle Code sections that match `data/statutes.json`
(CVC 21453 red light, CVC 21703 following too closely) plus CVC 21801 (left-turn yield)
and CVC 22350 (basic speed law) for the disputed case.
""", encoding="utf-8")


def main() -> None:
    BASE.mkdir(parents=True, exist_ok=True)
    case_pursue(BASE / "case_pursue_redlight")
    case_decline(BASE / "case_decline_rearend")
    case_escalate(BASE / "case_escalate_leftturn")
    write_readme(BASE)
    files = sorted(p for p in BASE.rglob("*") if p.is_file())
    print(f"Wrote {len(files)} files under {BASE}:")
    for p in files:
        print(f"  {p.relative_to(BASE)}  ({p.stat().st_size} bytes)")


if __name__ == "__main__":
    main()

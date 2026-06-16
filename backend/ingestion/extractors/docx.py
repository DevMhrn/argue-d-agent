"""DOCX extractor.

Uses python-docx to read paragraphs + tables. DOCX has no native page concept;
we split into logical pages at top-level Heading 1 boundaries, falling back to
one page if there are no headings.
"""
from __future__ import annotations

from .base import ExtractedDocument, ExtractedPage


class DocxExtractor:
    mime_types = (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )

    def extract(self, file_bytes: bytes, *, filename: str) -> ExtractedDocument:
        import io

        from docx import Document  # type: ignore[import-not-found]

        doc = Document(io.BytesIO(file_bytes))

        # Group paragraphs by Heading-1 boundary into logical pages.
        groups: list[list[str]] = [[]]
        for p in doc.paragraphs:
            style = (p.style.name or "") if p.style else ""
            if style.startswith("Heading 1") and groups[-1]:
                groups.append([])
            if p.text.strip():
                groups[-1].append(p.text)

        # Tables: append as text blocks at the end (simple for v1).
        for table in doc.tables:
            rows = []
            for row in table.rows:
                rows.append(" | ".join(cell.text for cell in row.cells))
            if rows:
                groups[-1].append("\n".join(rows))

        pages = [
            ExtractedPage(page_number=i + 1, text="\n".join(lines).strip())
            for i, lines in enumerate(groups)
            if lines
        ]
        if not pages:
            pages = [ExtractedPage(page_number=1, text="")]

        return ExtractedDocument(
            pages=pages,
            document_metadata={"extractor": "docx"},
        )
